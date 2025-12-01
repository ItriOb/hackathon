"""
BRD Processing Application
Complete backend system for processing Business Requirement Documents
using Streamlit, LangChain, AWS Bedrock Claude Sonnet 4, and python-docx
"""

import streamlit as st
import json
import logging
from typing import List, Dict, Optional, Any
from datetime import datetime
from pathlib import Path
import traceback

# Document processing
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# AWS and LangChain
import boto3
from langchain_aws import ChatBedrock
from langchain.prompts import PromptTemplate
from langchain.output_parsers import PydanticOutputParser

# Data models
from pydantic import BaseModel, Field, validator

# Excel export
import pandas as pd
from io import BytesIO

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ============================================================================
# PYDANTIC MODELS
# ============================================================================

class UserStory(BaseModel):
    """User Story model"""
    id: str = Field(description="Unique identifier for the user story")
    title: str = Field(description="Short title of the user story")
    description: str = Field(description="Full user story description")
    role: str = Field(description="User role (As a...)")
    action: str = Field(description="What the user wants (I want...)")
    benefit: str = Field(description="Why they want it (So that...)")
    acceptance_criteria: List[str] = Field(description="List of acceptance criteria")
    priority: str = Field(default="Medium", description="Priority: High, Medium, Low")
    
    @validator('priority')
    def validate_priority(cls, v):
        if v not in ['High', 'Medium', 'Low']:
            return 'Medium'
        return v


class Epic(BaseModel):
    """Epic model"""
    id: str = Field(description="Unique identifier for the epic")
    title: str = Field(description="Epic title")
    description: str = Field(description="Detailed description of the epic")
    business_value: str = Field(description="Business value and rationale")
    priority: str = Field(default="Medium", description="Priority: High, Medium, Low")
    user_stories: List[UserStory] = Field(default=[], description="List of user stories")
    
    @validator('priority')
    def validate_priority(cls, v):
        if v not in ['High', 'Medium', 'Low']:
            return 'Medium'
        return v


class DataModelField(BaseModel):
    """Data model field specification"""
    name: str = Field(description="Field name")
    type: str = Field(description="Data type")
    required: bool = Field(default=True, description="Whether field is required")
    description: str = Field(description="Field description")
    constraints: Optional[str] = Field(default=None, description="Validation constraints")


class DataModel(BaseModel):
    """Data model specification"""
    entity_name: str = Field(description="Name of the entity/model")
    description: str = Field(description="Description of the entity")
    fields: List[DataModelField] = Field(description="List of fields")
    relationships: List[str] = Field(default=[], description="Relationships with other entities")
    pydantic_code: str = Field(description="Generated Pydantic model code")


class TestCase(BaseModel):
    """Test case specification"""
    test_id: str = Field(description="Unique test identifier")
    test_name: str = Field(description="Descriptive test name")
    test_type: str = Field(description="Type: unit, integration, e2e")
    description: str = Field(description="What the test validates")
    test_code: str = Field(description="pytest test code")
    user_story_id: Optional[str] = Field(default=None, description="Related user story ID")


class GherkinScenario(BaseModel):
    """Gherkin BDD scenario"""
    feature: str = Field(description="Feature name")
    scenario: str = Field(description="Scenario name")
    given: List[str] = Field(description="Given steps (preconditions)")
    when: List[str] = Field(description="When steps (actions)")
    then: List[str] = Field(description="Then steps (expected outcomes)")
    user_story_id: Optional[str] = Field(default=None, description="Related user story ID")


class BRDProcessingResult(BaseModel):
    """Complete BRD processing result"""
    document_name: str
    processed_at: str
    epics: List[Epic]
    data_models: List[DataModel]
    test_cases: List[TestCase]
    gherkin_scenarios: List[GherkinScenario]


# ============================================================================
# DOCX PARSER
# ============================================================================

class DOCXParser:
    """Parse DOCX files and extract structured content"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document = Document(file_path)
        
    def extract_sections(self) -> Dict[str, str]:
        """Extract text organized by section headings"""
        sections = {}
        current_heading = "Introduction"
        current_content = []
        
        for para in self.document.paragraphs:
            if para.style.name.startswith('Heading'):
                # Save previous section
                if current_content:
                    sections[current_heading] = '\n'.join(current_content)
                
                # Start new section
                current_heading = para.text.strip()
                current_content = []
            else:
                if para.text.strip():
                    current_content.append(para.text.strip())
        
        # Save last section
        if current_content:
            sections[current_heading] = '\n'.join(current_content)
        
        logger.info(f"Extracted {len(sections)} sections from document")
        return sections
    
    def extract_tables(self) -> List[Dict[str, Any]]:
        """Extract tables from document"""
        tables_data = []
        
        for i, table in enumerate(self.document.tables):
            table_data = {
                'table_id': f"table_{i+1}",
                'rows': []
            }
            
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data['rows'].append(row_data)
            
            tables_data.append(table_data)
        
        logger.info(f"Extracted {len(tables_data)} tables from document")
        return tables_data
    
    def get_full_text(self) -> str:
        """Get complete document text"""
        return '\n\n'.join([para.text for para in self.document.paragraphs if para.text.strip()])
    
    def identify_section_type(self, section_text: str) -> str:
        """Classify section type based on content"""
        section_lower = section_text.lower()
        
        if any(keyword in section_lower for keyword in ['objective', 'goal', 'purpose']):
            return 'objectives'
        elif any(keyword in section_lower for keyword in ['functional requirement', 'feature', 'capability']):
            return 'functional_requirements'
        elif any(keyword in section_lower for keyword in ['non-functional', 'performance', 'security']):
            return 'non_functional_requirements'
        elif any(keyword in section_lower for keyword in ['user', 'role', 'actor', 'stakeholder']):
            return 'user_roles'
        elif any(keyword in section_lower for keyword in ['constraint', 'limitation', 'assumption']):
            return 'constraints'
        else:
            return 'general'


# ============================================================================
# AWS BEDROCK CLIENT
# ============================================================================

class BedrockClient:
    """AWS Bedrock client for Claude Sonnet 4"""
    
    def __init__(self, region_name: str = "us-east-1"):
        self.model_id = "anthropic.claude-sonnet-4-20250514"
        self.region_name = region_name
        
        try:
            self.client = ChatBedrock(
                model_id=self.model_id,
                region_name=self.region_name,
                model_kwargs={
                    "temperature": 0.3,
                    "top_p": 0.9,
                    "max_tokens": 4096
                }
            )
            logger.info(f"Initialized Bedrock client with model {self.model_id}")
        except Exception as e:
            logger.error(f"Failed to initialize Bedrock client: {e}")
            raise
    
    def generate_structured_output(
        self,
        prompt: str,
        response_model: BaseModel,
        system_prompt: Optional[str] = None
    ) -> BaseModel:
        """Generate structured output using Pydantic model"""
        try:
            # Create output parser
            parser = PydanticOutputParser(pydantic_object=response_model)
            
            # Add format instructions to prompt
            format_instructions = parser.get_format_instructions()
            full_prompt = f"{prompt}\n\n{format_instructions}"
            
            # Generate response
            messages = []
            if system_prompt:
                messages.append({"role": "system", "content": system_prompt})
            messages.append({"role": "user", "content": full_prompt})
            
            response = self.client.invoke(full_prompt)
            
            # Parse response
            parsed_output = parser.parse(response.content)
            
            logger.info(f"Successfully generated structured output for {response_model.__name__}")
            return parsed_output
            
        except Exception as e:
            logger.error(f"Error generating structured output: {e}")
            logger.error(traceback.format_exc())
            raise
    
    def generate_completion(self, prompt: str, system_prompt: Optional[str] = None) -> str:
        """Generate text completion"""
        try:
            full_prompt = prompt
            if system_prompt:
                full_prompt = f"{system_prompt}\n\n{prompt}"
            
            response = self.client.invoke(full_prompt)
            return response.content
            
        except Exception as e:
            logger.error(f"Error generating completion: {e}")
            raise


# ============================================================================
# PROMPT TEMPLATES
# ============================================================================

class PromptTemplates:
    """Collection of prompt templates for different generation tasks"""
    
    EPIC_GENERATION = """
Based on the following BRD section, generate epics with their associated user stories.

BRD Section:
{brd_section}

Generate 2-4 epics that cover the key functional areas described in this section.
Each epic should have:
- A clear title
- Detailed description
- Business value explanation
- Priority (High/Medium/Low)
- 3-5 user stories in the format: "As a [role], I want [action] so that [benefit]"

Each user story should include:
- Unique ID (format: US-XXX)
- Title
- Full description
- Role, action, and benefit clearly separated
- 3-5 specific acceptance criteria
- Priority

Provide the output as a JSON array of epic objects.
"""

    USER_STORY_EXPANSION = """
Based on the following epic, generate detailed user stories.

Epic: {epic_title}
Description: {epic_description}

Generate 5-8 comprehensive user stories that break down this epic into implementable units.
Each user story must follow the format:
- As a [specific role]
- I want [specific action]
- So that [clear benefit]

Include:
- Unique ID (format: US-XXX)
- Descriptive title
- Full description
- 3-5 testable acceptance criteria
- Priority based on business value

Provide the output as a JSON array of user story objects.
"""

    DATA_MODEL_GENERATION = """
Based on the following requirements, identify and design data models.

Requirements:
{requirements_text}

Analyze the requirements and:
1. Identify key entities/objects
2. Define fields for each entity with appropriate data types
3. Specify validation constraints
4. Define relationships between entities
5. Generate complete Pydantic model code

For each data model, provide:
- Entity name
- Description
- List of fields with: name, type, required flag, description, constraints
- Relationships with other entities
- Complete Pydantic model code implementation

Provide the output as a JSON array of data model objects.
"""

    PYTEST_GENERATION = """
Based on the following user story, generate comprehensive pytest test cases.

User Story: {user_story}

Generate test cases covering:
1. Happy path scenarios
2. Edge cases
3. Error handling
4. Boundary conditions
5. Input validation

For each test case, provide:
- Unique test ID (format: TC-XXX)
- Descriptive test name (test_xxx_xxx format)
- Test type (unit/integration/e2e)
- Description of what is being tested
- Complete pytest test code with:
  - Proper fixtures if needed
  - Arrange-Act-Assert structure
  - Clear assertions
  - Good documentation

Provide the output as a JSON array of test case objects.
"""

    GHERKIN_GENERATION = """
Based on the following user story and acceptance criteria, generate Gherkin BDD scenarios.

User Story: {user_story}
Acceptance Criteria:
{acceptance_criteria}

Generate Gherkin scenarios in Given-When-Then format.
Create multiple scenarios covering:
1. Main success scenario
2. Alternative flows
3. Error scenarios
4. Edge cases

For each scenario, provide:
- Feature name
- Scenario name
- Given steps (preconditions)
- When steps (actions)
- Then steps (expected outcomes)

Ensure steps are:
- Specific and testable
- Written from user perspective
- Clear and unambiguous

Provide the output as a JSON array of Gherkin scenario objects.
"""


# ============================================================================
# BRD PROCESSING PIPELINE
# ============================================================================

class BRDPipeline:
    """Main pipeline for processing BRD documents"""
    
    def __init__(self, bedrock_client: BedrockClient):
        self.bedrock = bedrock_client
        self.templates = PromptTemplates()
        
    def generate_epics(self, brd_sections: Dict[str, str], progress_callback=None) -> List[Epic]:
        """Generate epics from BRD sections"""
        epics = []
        
        # Filter relevant sections
        relevant_sections = {
            k: v for k, v in brd_sections.items()
            if any(keyword in k.lower() for keyword in ['requirement', 'feature', 'functional', 'objective'])
        }
        
        total_sections = len(relevant_sections)
        
        for idx, (section_name, section_content) in enumerate(relevant_sections.items()):
            try:
                if progress_callback:
                    progress_callback(f"Generating epics from section: {section_name}", idx / total_sections)
                
                # Generate prompt
                prompt = self.templates.EPIC_GENERATION.format(brd_section=section_content)
                
                # Generate completion
                response = self.bedrock.generate_completion(prompt)
                
                # Parse JSON response
                # In production, use structured output parser
                # For now, extract JSON from response
                response_json = self._extract_json(response)
                
                if response_json and isinstance(response_json, list):
                    for epic_data in response_json:
                        epic = Epic(**epic_data)
                        epics.append(epic)
                
                logger.info(f"Generated {len(epics)} epics from section: {section_name}")
                
            except Exception as e:
                logger.error(f"Error generating epics from section {section_name}: {e}")
                continue
        
        return epics
    
    def generate_user_stories(self, epics: List[Epic], progress_callback=None) -> List[Epic]:
        """Expand epics with more detailed user stories"""
        total_epics = len(epics)
        
        for idx, epic in enumerate(epics):
            try:
                if progress_callback:
                    progress_callback(f"Expanding user stories for epic: {epic.title}", idx / total_epics)
                
                # If epic already has stories, skip
                if len(epic.user_stories) >= 5:
                    continue
                
                prompt = self.templates.USER_STORY_EXPANSION.format(
                    epic_title=epic.title,
                    epic_description=epic.description
                )
                
                response = self.bedrock.generate_completion(prompt)
                response_json = self._extract_json(response)
                
                if response_json and isinstance(response_json, list):
                    user_stories = [UserStory(**story_data) for story_data in response_json]
                    epic.user_stories.extend(user_stories)
                
                logger.info(f"Generated {len(epic.user_stories)} user stories for epic: {epic.title}")
                
            except Exception as e:
                logger.error(f"Error generating user stories for epic {epic.title}: {e}")
                continue
        
        return epics
    
    def generate_data_models(self, brd_sections: Dict[str, str], progress_callback=None) -> List[DataModel]:
        """Generate data models from requirements"""
        data_models = []
        
        # Combine relevant sections
        requirements_text = "\n\n".join([
            f"{k}:\n{v}" for k, v in brd_sections.items()
            if any(keyword in k.lower() for keyword in ['requirement', 'data', 'entity', 'model'])
        ])
        
        try:
            if progress_callback:
                progress_callback("Generating data models from requirements", 0.5)
            
            prompt = self.templates.DATA_MODEL_GENERATION.format(requirements_text=requirements_text)
            
            response = self.bedrock.generate_completion(prompt)
            response_json = self._extract_json(response)
            
            if response_json and isinstance(response_json, list):
                for model_data in response_json:
                    data_model = DataModel(**model_data)
                    data_models.append(data_model)
            
            logger.info(f"Generated {len(data_models)} data models")
            
        except Exception as e:
            logger.error(f"Error generating data models: {e}")
        
        return data_models
    
    def generate_test_cases(self, epics: List[Epic], progress_callback=None) -> List[TestCase]:
        """Generate pytest test cases from user stories"""
        test_cases = []
        
        all_stories = []
        for epic in epics:
            all_stories.extend(epic.user_stories)
        
        total_stories = len(all_stories)
        
        for idx, story in enumerate(all_stories):
            try:
                if progress_callback:
                    progress_callback(f"Generating tests for: {story.title}", idx / total_stories)
                
                user_story_text = f"{story.description}\nAcceptance Criteria:\n" + \
                                "\n".join([f"- {ac}" for ac in story.acceptance_criteria])
                
                prompt = self.templates.PYTEST_GENERATION.format(user_story=user_story_text)
                
                response = self.bedrock.generate_completion(prompt)
                response_json = self._extract_json(response)
                
                if response_json and isinstance(response_json, list):
                    for test_data in response_json:
                        test_data['user_story_id'] = story.id
                        test_case = TestCase(**test_data)
                        test_cases.append(test_case)
                
            except Exception as e:
                logger.error(f"Error generating tests for story {story.id}: {e}")
                continue
        
        logger.info(f"Generated {len(test_cases)} test cases")
        return test_cases
    
    def generate_gherkin_scenarios(self, epics: List[Epic], progress_callback=None) -> List[GherkinScenario]:
        """Generate Gherkin BDD scenarios from user stories"""
        scenarios = []
        
        all_stories = []
        for epic in epics:
            all_stories.extend(epic.user_stories)
        
        total_stories = len(all_stories)
        
        for idx, story in enumerate(all_stories):
            try:
                if progress_callback:
                    progress_callback(f"Generating Gherkin for: {story.title}", idx / total_stories)
                
                user_story_text = story.description
                acceptance_criteria = "\n".join([f"- {ac}" for ac in story.acceptance_criteria])
                
                prompt = self.templates.GHERKIN_GENERATION.format(
                    user_story=user_story_text,
                    acceptance_criteria=acceptance_criteria
                )
                
                response = self.bedrock.generate_completion(prompt)
                response_json = self._extract_json(response)
                
                if response_json and isinstance(response_json, list):
                    for scenario_data in response_json:
                        scenario_data['user_story_id'] = story.id
                        scenario = GherkinScenario(**scenario_data)
                        scenarios.append(scenario)
                
            except Exception as e:
                logger.error(f"Error generating Gherkin for story {story.id}: {e}")
                continue
        
        logger.info(f"Generated {len(scenarios)} Gherkin scenarios")
        return scenarios
    
    def process_document(
        self,
        file_path: str,
        progress_callback=None
    ) -> BRDProcessingResult:
        """Process complete BRD document through pipeline"""
        try:
            # Parse document
            if progress_callback:
                progress_callback("Parsing DOCX document", 0.1)
            
            parser = DOCXParser(file_path)
            sections = parser.extract_sections()
            
            # Generate epics
            if progress_callback:
                progress_callback("Generating epics", 0.2)
            epics = self.generate_epics(sections, progress_callback)
            
            # Generate user stories
            if progress_callback:
                progress_callback("Generating user stories", 0.4)
            epics = self.generate_user_stories(epics, progress_callback)
            
            # Generate data models
            if progress_callback:
                progress_callback("Generating data models", 0.6)
            data_models = self.generate_data_models(sections, progress_callback)
            
            # Generate test cases
            if progress_callback:
                progress_callback("Generating test cases", 0.7)
            test_cases = self.generate_test_cases(epics, progress_callback)
            
            # Generate Gherkin scenarios
            if progress_callback:
                progress_callback("Generating Gherkin scenarios", 0.9)
            gherkin_scenarios = self.generate_gherkin_scenarios(epics, progress_callback)
            
            # Create result
            result = BRDProcessingResult(
                document_name=Path(file_path).name,
                processed_at=datetime.now().isoformat(),
                epics=epics,
                data_models=data_models,
                test_cases=test_cases,
                gherkin_scenarios=gherkin_scenarios
            )
            
            if progress_callback:
                progress_callback("Processing complete", 1.0)
            
            logger.info("Document processing completed successfully")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            logger.error(traceback.format_exc())
            raise
    
    def _extract_json(self, response: str) -> Optional[Any]:
        """Extract JSON from LLM response"""
        try:
            # Try to find JSON array or object in response
            import re
            
            # Look for JSON array
            json_match = re.search(r'\[[\s\S]*\]', response)
            if json_match:
                return json.loads(json_match.group())
            
            # Look for JSON object
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                return json.loads(json_match.group())
            
            # Try parsing entire response
            return json.loads(response)
            
        except Exception as e:
            logger.warning(f"Could not extract JSON from response: {e}")
            return None


# ============================================================================
# EXPORT UTILITIES
# ============================================================================

class Exporters:
    """Export utilities for generated artifacts"""
    
    @staticmethod
    def export_to_json(result: BRDProcessingResult, file_path: str):
        """Export results to JSON file"""
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(result.dict(), f, indent=2, ensure_ascii=False)
            logger.info(f"Exported results to JSON: {file_path}")
        except Exception as e:
            logger.error(f"Error exporting to JSON: {e}")
            raise
    
    @staticmethod
    def export_to_excel(result: BRDProcessingResult) -> BytesIO:
        """Export results to Excel with multiple sheets"""
        try:
            output = BytesIO()
            
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                # Epics sheet
                epics_data = []
                for epic in result.epics:
                    epics_data.append({
                        'Epic ID': epic.id,
                        'Title': epic.title,
                        'Description': epic.description,
                        'Business Value': epic.business_value,
                        'Priority': epic.priority,
                        'User Stories Count': len(epic.user_stories)
                    })
                df_epics = pd.DataFrame(epics_data)
                df_epics.to_excel(writer, sheet_name='Epics', index=False)
                
                # User Stories sheet
                stories_data = []
                for epic in result.epics:
                    for story in epic.user_stories:
                        stories_data.append({
                            'Story ID': story.id,
                            'Epic ID': epic.id,
                            'Title': story.title,
                            'Role': story.role,
                            'Action': story.action,
                            'Benefit': story.benefit,
                            'Priority': story.priority,
                            'Acceptance Criteria': '\n'.join(story.acceptance_criteria)
                        })
                df_stories = pd.DataFrame(stories_data)
                df_stories.to_excel(writer, sheet_name='User Stories', index=False)
                
                # Data Models sheet
                models_data = []
                for model in result.data_models:
                    models_data.append({
                        'Entity Name': model.entity_name,
                        'Description': model.description,
                        'Fields Count': len(model.fields),
                        'Relationships': ', '.join(model.relationships)
                    })
                df_models = pd.DataFrame(models_data)
                df_models.to_excel(writer, sheet_name='Data Models', index=False)
                
                # Test Cases sheet
                tests_data = []
                for test in result.test_cases:
                    tests_data.append({
                        'Test ID': test.test_id,
                        'Test Name': test.test_name,
                        'Type': test.test_type,
                        'Description': test.description,
                        'User Story ID': test.user_story_id
                    })
                df_tests = pd.DataFrame(tests_data)
                df_tests.to_excel(writer, sheet_name='Test Cases', index=False)
                
                # Gherkin Scenarios sheet
                gherkin_data = []
                for scenario in result.gherkin_scenarios:
                    gherkin_data.append({
                        'Feature': scenario.feature,
                        'Scenario': scenario.scenario,
                        'Given': '\n'.join(scenario.given),
                        'When': '\n'.join(scenario.when),
                        'Then': '\n'.join(scenario.then),
                        'User Story ID': scenario.user_story_id
                    })
                df_gherkin = pd.DataFrame(gherkin_data)
                df_gherkin.to_excel(writer, sheet_name='Gherkin Scenarios', index=False)
            
            output.seek(0)
            logger.info("Exported results to Excel")
            return output
            
        except Exception as e:
            logger.error(f"Error exporting to Excel: {e}")
            raise
    
    @staticmethod
    def generate_traceability_matrix(result: BRDProcessingResult) -> pd.DataFrame:
        """Generate traceability matrix linking requirements to stories to tests"""
        matrix_data = []
        
        for epic in result.epics:
            for story in epic.user_stories:
                # Find related tests
                related_tests = [
                    test for test in result.test_cases
                    if test.user_story_id == story.id
                ]
                
                # Find related Gherkin scenarios
                related_scenarios = [
                    scenario for scenario in result.gherkin_scenarios
                    if scenario.user_story_id == story.id
                ]
                
                matrix_data.append({
                    'Epic ID': epic.id,
                    'Epic Title': epic.title,
                    'Story ID': story.id,
                    'Story Title': story.title,
                    'Test Cases Count': len(related_tests),
                    'Test Case IDs': ', '.join([test.test_id for test in related_tests]),
                    'Gherkin Scenarios Count': len(related_scenarios),
                    'Coverage': 'Full' if related_tests and related_scenarios else 'Partial'
                })
        
        return pd.DataFrame(matrix_data)


# ============================================================================
# STREAMLIT APPLICATION
# ============================================================================

def initialize_session_state():
    """Initialize Streamlit session state"""
    if 'processed_result' not in st.session_state:
        st.session_state.processed_result = None
    if 'bedrock_client' not in st.session_state:
        st.session_state.bedrock_client = None
    if 'processing_status' not in st.session_state:
        st.session_state.processing_status = ""


def main():
    """Main Streamlit application"""
    st.set_page_config(
        page_title="BRD Processing System",
        page_icon="📄",
        layout="wide"
    )
    
    initialize_session_state()
    
    st.title("📄 BRD Processing System")
    st.markdown("""
    Upload a Business Requirements Document (DOCX) to automatically generate:
    - **Epics & User Stories**
    - **Data Models**
    - **Functional Tests (pytest)**
    - **Gherkin BDD Scenarios**
    """)
    
    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        aws_region = st.selectbox(
            "AWS Region",
            ["us-east-1", "us-west-2", "eu-west-1"],
            index=0
        )
        
        if st.button("Initialize AWS Bedrock", type="primary"):
            try:
                with st.spinner("Connecting to AWS Bedrock..."):
                    st.session_state.bedrock_client = BedrockClient(region_name=aws_region)
                st.success("✅ Connected to AWS Bedrock")
            except Exception as e:
                st.error(f"❌ Failed to connect: {str(e)}")
        
        st.divider()
        
        st.header("📊 Processing Options")
        generate_epics = st.checkbox("Generate Epics & User Stories", value=True)
        generate_models = st.checkbox("Generate Data Models", value=True)
        generate_tests = st.checkbox("Generate Test Cases", value=True)
        generate_gherkin = st.checkbox("Generate Gherkin Scenarios", value=True)
    
    # Main content area
    uploaded_file = st.file_uploader(
        "Upload BRD Document (DOCX)",
        type=['docx'],
        help="Upload your Business Requirements Document in DOCX format"
    )
    
    if uploaded_file is not None:
        # Save uploaded file
        temp_path = f"/tmp/{uploaded_file.name}"
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        st.success(f"✅ Uploaded: {uploaded_file.name}")
        
        # Document preview
        with st.expander("📖 Document Preview"):
            try:
                parser = DOCXParser(temp_path)
                sections = parser.extract_sections()
                
                st.write(f"**Sections found:** {len(sections)}")
                for section_name in sections.keys():
                    st.write(f"- {section_name}")
            except Exception as e:
                st.error(f"Error parsing document: {e}")
        
        # Process button
        if st.button("🚀 Process Document", type="primary", disabled=st.session_state.bedrock_client is None):
            if st.session_state.bedrock_client is None:
                st.error("Please initialize AWS Bedrock first (see sidebar)")
            else:
                # Progress tracking
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                def update_progress(message, progress):
                    status_text.text(message)
                    progress_bar.progress(progress)
                
                try:
                    # Create pipeline
                    pipeline = BRDPipeline(st.session_state.bedrock_client)
                    
                    # Process document
                    with st.spinner("Processing document..."):
                        result = pipeline.process_document(
                            temp_path,
                            progress_callback=update_progress
                        )
                    
                    st.session_state.processed_result = result
                    st.success("✅ Processing complete!")
                    
                except Exception as e:
                    st.error(f"❌ Processing failed: {str(e)}")
                    logger.error(traceback.format_exc())
    
    # Display results
    if st.session_state.processed_result:
        result = st.session_state.processed_result
        
        st.divider()
        st.header("📊 Processing Results")
        
        # Summary metrics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Epics", len(result.epics))
        with col2:
            total_stories = sum(len(epic.user_stories) for epic in result.epics)
            st.metric("User Stories", total_stories)
        with col3:
            st.metric("Data Models", len(result.data_models))
        with col4:
            st.metric("Test Cases", len(result.test_cases))
        
        # Tabs for different outputs
        tab1, tab2, tab3, tab4, tab5 = st.tabs([
            "📋 Epics & Stories",
            "🗂️ Data Models",
            "🧪 Test Cases",
            "🥒 Gherkin",
            "📥 Export"
        ])
        
        with tab1:
            st.subheader("Epics & User Stories")
            for epic in result.epics:
                with st.expander(f"🎯 {epic.title} ({epic.priority} Priority)"):
                    st.write(f"**ID:** {epic.id}")
                    st.write(f"**Description:** {epic.description}")
                    st.write(f"**Business Value:** {epic.business_value}")
                    
                    st.write("**User Stories:**")
                    for story in epic.user_stories:
                        st.write(f"- **{story.id}**: {story.title}")
                        st.write(f"  - Role: {story.role}")
                        st.write(f"  - Action: {story.action}")
                        st.write(f"  - Benefit: {story.benefit}")
                        st.write(f"  - Priority: {story.priority}")
                        st.write("  - Acceptance Criteria:")
                        for ac in story.acceptance_criteria:
                            st.write(f"    - {ac}")
        
        with tab2:
            st.subheader("Data Models")
            for model in result.data_models:
                with st.expander(f"📊 {model.entity_name}"):
                    st.write(f"**Description:** {model.description}")
                    
                    st.write("**Fields:**")
                    fields_df = pd.DataFrame([
                        {
                            'Name': f.name,
                            'Type': f.type,
                            'Required': f.required,
                            'Description': f.description
                        }
                        for f in model.fields
                    ])
                    st.dataframe(fields_df, use_container_width=True)
                    
                    if model.relationships:
                        st.write("**Relationships:**")
                        for rel in model.relationships:
                            st.write(f"- {rel}")
                    
                    st.code(model.pydantic_code, language='python')
        
        with tab3:
            st.subheader("Test Cases")
            for test in result.test_cases:
                with st.expander(f"🧪 {test.test_name}"):
                    st.write(f"**ID:** {test.test_id}")
                    st.write(f"**Type:** {test.test_type}")
                    st.write(f"**Description:** {test.description}")
                    if test.user_story_id:
                        st.write(f"**User Story:** {test.user_story_id}")
                    st.code(test.test_code, language='python')
        
        with tab4:
            st.subheader("Gherkin Scenarios")
            for scenario in result.gherkin_scenarios:
                with st.expander(f"🥒 {scenario.feature} - {scenario.scenario}"):
                    if scenario.user_story_id:
                        st.write(f"**User Story:** {scenario.user_story_id}")
                    
                    st.write("**Given:**")
                    for step in scenario.given:
                        st.write(f"- {step}")
                    
                    st.write("**When:**")
                    for step in scenario.when:
                        st.write(f"- {step}")
                    
                    st.write("**Then:**")
                    for step in scenario.then:
                        st.write(f"- {step}")
        
        with tab5:
            st.subheader("Export Results")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                # Export to JSON
                json_data = json.dumps(result.dict(), indent=2)
                st.download_button(
                    label="📄 Download JSON",
                    data=json_data,
                    file_name=f"brd_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json"
                )
            
            with col2:
                # Export to Excel
                try:
                    excel_buffer = Exporters.export_to_excel(result)
                    st.download_button(
                        label="📊 Download Excel",
                        data=excel_buffer,
                        file_name=f"brd_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                except Exception as e:
                    st.error(f"Error generating Excel: {e}")
            
            with col3:
                # Traceability Matrix
                try:
                    matrix_df = Exporters.generate_traceability_matrix(result)
                    csv_data = matrix_df.to_csv(index=False)
                    st.download_button(
                        label="🔗 Download Traceability Matrix",
                        data=csv_data,
                        file_name=f"traceability_matrix_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                        mime="text/csv"
                    )
                except Exception as e:
                    st.error(f"Error generating matrix: {e}")


if __name__ == "__main__":
    main()
